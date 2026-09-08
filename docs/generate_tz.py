#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Генерация технического задания (ГОСТ 19.201-78 / ГОСТ 19.106-78)."""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor, Twips
from docx.enum.section import WD_ORIENT

FONT = "Times New Roman"
DOC_CODE = ""

STUDENTS = [
    {
        "full": "Богачев Егор Андреевич",
        "short": "Е. А. Богачев",
        "initials": "Богачев Е. А.",
        "code": "ИСИП.241П.005-01 ТЗ",
        "out": "/workspace/docs/TZ_Bogachev_ISIP-24-1P.docx",
        "role": "студент группы ИСИП-24-1П",
        "role_long": "студент группы ИСИП-24-1П",
    },
    {
        "full": "Бальшинов Добдан Баирович",
        "short": "Д. Б. Бальшинов",
        "initials": "Бальшинов Д. Б.",
        "code": "ИСИП.241П.006-01 ТЗ",
        "out": "/workspace/docs/TZ_Balshinov_ISIP-24-1P.docx",
        "role": "студент группы ИСИП-24-1П",
        "role_long": "студент группы ИСИП-24-1П",
    },
]

def set_run_font(run, size=14, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)


def set_paragraph_font(paragraph, size=14):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=run.bold, italic=run.italic)


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_run_font(run, size=14)
    r = run._r
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r.append(fc1)
    r.append(instr)
    r.append(fc2)


def disable_widow_control(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    widow = OxmlElement("w:widowControl")
    widow.set(qn("w:val"), "false")
    pPr.append(widow)


def configure_normal_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(14)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def setup_section(section, different_first=True):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Mm(30)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(10)
    section.different_first_page_header_footer = different_first

    # ГОСТ 19.201-78: номера листов — в верхней части над текстом
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.first_line_indent = Cm(0)
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    add_page_field(hp)

    first = section.first_page_header
    first.is_linked_to_previous = False
    fp = first.paragraphs[0]
    fp.text = ""

    footer = section.footer
    footer.is_linked_to_previous = False
    fpar = footer.paragraphs[0]
    fpar.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fpar.paragraph_format.first_line_indent = Cm(0)
    fpar.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = fpar.add_run(DOC_CODE)
    set_run_font(run, size=10, italic=True)


def add_empty(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run("")
        set_run_font(run, size=14)


def add_center(doc, text, size=14, bold=False, italic=False, space_before=0, space_after=0, caps=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text.upper() if caps else text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_body(doc, text, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25) if first_indent else Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=14)
    return p


def add_section_heading(doc, text):
    """ГОСТ 19.106-78: заголовки разделов — прописными, симметрично тексту."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    set_run_font(run, size=14, bold=True)
    return p


def add_subheading(doc, text):
    """ГОСТ 19.106-78: заголовки подразделов — с абзаца, с прописной."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True)
    return p


def set_cell_shading(cell, fill="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def shade_header_row(row):
    for cell in row.cells:
        set_cell_shading(cell, "E7E6E6")


def fill_cell(cell, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, center_v=True):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    set_cell_border(cell)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center" if center_v else "top")
    tcPr.append(vAlign)


def add_table(doc, headers, rows, col_widths=None, header_size=11, cell_size=11):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)

    # ширина таблицы = 170 мм (210 − 30 − 10)
    usable = Cm(17.0)
    if col_widths:
        total = sum(col_widths)
        widths = [usable * (w / total) for w in col_widths]
    else:
        widths = [usable / len(headers)] * len(headers)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        fill_cell(cell, h, size=header_size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_header_row(table.rows[0])

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            fill_cell(table.rows[r_idx + 1].cells[c_idx], val, size=cell_size, align=align)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
    return table


def add_caption(doc, text):
    """ГОСТ 19.106-78: название таблицы помещают над таблицей, слева."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=14, italic=False)
    return p


def build_title_page(doc, author):
    add_center(
        doc,
        "МИНИСТЕРСТВО ПРОСВЕЩЕНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        size=14,
        bold=True,
        space_after=0,
    )
    add_center(
        doc,
        "Образовательная организация среднего профессионального образования",
        size=14,
        italic=True,
        space_after=0,
    )
    add_center(
        doc,
        "Специальность 09.02.07 «Информационные системы и программирование»",
        size=14,
        space_after=0,
    )
    add_center(doc, "Группа ИСИП-24-1П", size=14, bold=True, space_after=6)

    add_empty(doc, 1)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left, right = table.rows[0].cells

    def stamp_block(cell, title):
        cell.text = ""
        lines = [
            (title, True),
            ("Преподаватель", False) if title == "СОГЛАСОВАНО" else ("Председатель ЦК", False),
            ("_________________ / __________ /", False),
            ("«___» ______________ 2026 г.", False),
        ]
        if title == "УТВЕРЖДАЮ":
            lines[1] = ("Председатель цикловой комиссии", False)
        first = True
        for text, bold in lines:
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(text)
            set_run_font(run, size=12, bold=bold)

    stamp_block(left, "СОГЛАСОВАНО")
    stamp_block(right, "УТВЕРЖДАЮ")
    for cell in table.rows[0].cells:
        cell.width = Cm(8.5)

    add_empty(doc, 2)

    add_center(doc, "ТЕХНИЧЕСКОЕ ЗАДАНИЕ", size=16, bold=True, space_before=6, space_after=0)
    add_center(doc, "на разработку программы", size=14, space_after=0)
    add_center(
        doc,
        "«Магазин товаров для взрослых»",
        size=14,
        bold=True,
        space_after=6,
    )
    add_center(doc, "(учебный программный документ)", size=14, italic=True, space_after=12)

    add_center(doc, author["code"], size=14, bold=True, space_after=0)
    add_center(doc, "Листов 7", size=14, space_after=6)

    add_empty(doc, 1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for line in (
        "Исполнитель:",
        author["role"],
        author["full"],
        "",
        f"_________________ / {author['short']} /",
    ):
        if p.runs:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(line)
        set_run_font(run, size=14)

    add_empty(doc, 2)
    add_center(doc, "2026", size=14, bold=True)


def build_body(doc, author):
    add_section_heading(doc, "Содержание")
    for item in (
        "1. Введение",
        "2. Основания для разработки",
        "3. Назначение разработки",
        "4. Требования к программе",
        "5. Требования к программной документации",
        "6. Технико-экономические показатели",
        "7. Стадии и этапы разработки",
        "8. Порядок контроля и приёмки",
        "9. Приложение. Перечень терминов и документов",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(item)
        set_run_font(run, size=14)

    add_section_heading(doc, "1. Введение")
    add_body(
        doc,
        "Наименование программы: «Магазин товаров для взрослых» "
        "(краткое наименование — МТВ). Обозначение документа: "
        f"{author['code']}. Программа предназначена для учёта "
        "ассортимента, остатков и продаж розничного магазина товаров "
        "категории 18+. Область применения — розничная торговля "
        "бельём, средствами ухода, аксессуарами и сувенирной продукцией "
        "с ограничением по возрасту покупателя. Объект применения — "
        f"рабочее место продавца ({author['role_long']}).",
    )

    add_section_heading(doc, "2. Основания для разработки")
    add_body(
        doc,
        "Основанием является учебное задание по материалам лекции 2 "
        "«Общие требования к программным документам. Техническое задание», "
        "утверждённое преподавателем цикловой комиссии специальности 09.02.07 "
        "в сентябре 2026 года. Исполнитель — "
        f"{author['role_long']} {author['full']}. Наименование темы: "
        "«Разработка программы «Магазин товаров для взрослых»». "
        "Шифр темы: МТВ-2026.",
    )

    add_section_heading(doc, "3. Назначение разработки")
    add_subheading(doc, "3.1. Функциональное назначение")
    add_body(
        doc,
        "Программа обеспечивает ведение каталога товаров 18+, учёт "
        "остатков, оформление продаж и заказов только совершеннолетним "
        "покупателям и формирование отчётов по наличию и выручке.",
    )
    add_subheading(doc, "3.2. Эксплуатационное назначение")
    add_body(
        doc,
        "Программа эксплуатируется на ПЭВМ под управлением Windows 10/11 "
        "в однопользовательском режиме на кассовом месте магазина. "
        "Постоянное подключение к сети Интернет не требуется. "
        "Данные хранятся локально в файле базы.",
    )

    add_section_heading(doc, "4. Требования к программе")
    add_subheading(doc, "4.1. Требования к функциональным характеристикам")
    add_body(
        doc,
        "Программа должна выполнять следующие функции: подтверждение "
        "возраста покупателя (не моложе 18 лет) перед каждой продажей "
        "с обязательной отметкой кассира; ведение каталога товаров "
        "(название, артикул, категория: бельё, уход и косметика, "
        "аксессуары, сувениры; цена; остаток; признак 18+); добавление, "
        "изменение и списание товара; оформление продажи и заказа "
        "(дата, состав, способ получения: самовывоз или доставка "
        "в закрытой упаковке, статус, сумма) без хранения паспортных "
        "данных покупателя; поиск по названию и артикулу; фильтр "
        "по категории и наличию; сортировка по названию, цене и остатку; "
        "просмотр карточки товара; формирование отчётов «Остатки», "
        "«Продажи за период» и «Товарный чек» (без детализации позиций "
        "на чеке, выдаваемом покупателю, по запросу кассира — обезличенный "
        "вид «товар категории 18+») с выводом на экран, печать и "
        "сохранение в PDF; предупреждение о низком остатке; резервное "
        "копирование файла базы данных.",
    )
    add_body(
        doc,
        "Входные данные вводятся через формы интерфейса. Выходные данные — "
        "каталог, список продаж, печатные формы формата A4 и файл "
        "базы SQLite. При объёме до 400 товаров и 300 продаж время "
        "открытия каталога — не более 3 с, оформления продажи — не более "
        "3 с, формирования отчёта — не более 5 с.",
    )

    add_subheading(doc, "4.2. Требования к надёжности")
    add_body(
        doc,
        "Программа не должна аварийно завершаться при некорректном вводе: "
        "ошибка сообщается пользователю. Контролируются уникальность "
        "артикула, обязательность названия и цены, цена больше нуля, "
        "неотрицательный остаток, запрет продажи без отметки о возрасте "
        "18+ и запрет продажи сверх остатка. При оформлении продажи "
        "остаток уменьшается автоматически. Возврат в день продажи "
        "возвращает количество на склад. После сбоя восстановление — "
        "повторным запуском; при повреждении базы предлагается копия. "
        "Время восстановления при наличии копии — не более 10 минут.",
    )

    add_subheading(doc, "4.3. Условия эксплуатации")
    add_body(
        doc,
        "Условия эксплуатации соответствуют требованиям к ПЭВМ: температура "
        "от плюс 10 до плюс 35 °С, относительная влажность 40–80 %. "
        "Обслуживание: корректное завершение работы и резервная копия "
        "не реже одного раза в неделю. Персонал — кассир не моложе 18 лет "
        "после инструктажа до 30 минут. Работа с ПЭВМ — не ниже II группы "
        "по электробезопасности.",
    )

    add_subheading(doc, "4.4. Требования к составу и параметрам технических средств")
    add_body(
        doc,
        "Минимальный состав рабочего места: ПЭВМ с процессором 1,6 ГГц "
        "(два ядра), ОЗУ 4 Гбайт, 500 Мбайт свободного места, дисплей "
        "1366×768, клавиатура и мышь. Для печати — принтер либо сохранение "
        "в PDF. Для копий — USB-накопитель от 1 Гбайт.",
    )

    add_subheading(doc, "4.5. Требования к информационной и программной совместимости")
    add_body(
        doc,
        "Данные хранятся в файле SQLite. Основные сущности: товар, "
        "категория, продажа, позиция продажи. Паспортные и иные "
        "персональные данные покупателя не хранятся. Язык реализации — "
        "C# (.NET 8). Среда — Windows 10/11 x64 и .NET 8 Desktop Runtime. "
        "Импорт товаров допускается из CSV (UTF-8): артикул, название, "
        "категория, цена, остаток.",
    )

    add_subheading(doc, "4.6. Требования к маркировке и упаковке")
    add_body(
        doc,
        "На носителе указываются наименование программы, обозначение "
        f"{author['code'].replace(' ТЗ', '')}, версия, исполнитель и группа. "
        "Передача — на USB-накопителе со сшитым комплектом документации "
        "либо архивом ZIP по согласованному учебному каналу.",
    )

    add_subheading(doc, "4.7. Требования к транспортированию и хранению")
    add_body(
        doc,
        "Носитель перевозят в закрытой таре при температуре от плюс 5 "
        "до плюс 40 °С без воздействия осадков. Хранение — в отапливаемом "
        "помещении не менее срока хранения учебной работы, установленного "
        "образовательной организацией.",
    )

    add_subheading(doc, "4.8. Специальные требования")
    add_body(
        doc,
        "Язык интерфейса — русский. Основной экран — список товаров "
        "(артикул, название, категория, цена, остаток). Продажа "
        "недоступна, пока кассир не подтвердит возраст покупателя 18+. "
        "Товары с остатком ниже минимума выделяются красным. Экран "
        "кассы не должен быть обращён к покупателю с детальным "
        "наименованием позиций (режим «приватный чек»). Сохранение — "
        "кнопка «Сохранить» и Ctrl+S. Требования размещения в сети "
        "Интернет не предъявляются (учебная программа).",
    )

    add_section_heading(doc, "5. Требования к программной документации")
    add_body(
        doc,
        "Комплект документов по ГОСТ 19.101-77, оформление по ГОСТ 19.106-78: "
        "техническое задание; текст программы (12); описание программы (13); "
        "руководство оператора (34); программа и методика испытаний (51); "
        "пояснительная записка (81). Документы — формат A4, шрифт Times New Roman "
        "14 пт (таблицы 10–12 пт), интервал 1,5, выравнивание по ширине, "
        "абзацный отступ 1,25 см; поля: левое 30 мм, правое 10 мм, верхнее "
        "и нижнее по 20 мм. Номера листов — вверху над текстом.",
    )

    add_section_heading(doc, "6. Технико-экономические показатели")
    add_body(
        doc,
        "Разработка учебная, стоимость лицензии для заказчика — 0 руб. "
        "Эффект — учёт остатков и продаж без бумажного журнала при "
        "соблюдении ограничения 18+. Потребность — один рабочий экземпляр "
        "на кассу и один контрольный для преподавателя. Преимущество "
        "перед универсальными кассовыми облаками — локальная работа "
        "и режим приватного чека; недостаток — нет эквайринга и "
        "онлайн-витрины, что для учебного изделия допустимо.",
    )

    add_section_heading(doc, "7. Стадии и этапы разработки")
    add_body(
        doc,
        "Стадии по ГОСТ 19.102-77: техническое задание; рабочий проект "
        "(совмещён с техническим); внедрение. Эскизный проект не выделяется. "
        f"Исполнитель всех стадий — {author['initials']}, группа ИСИП-24-1П.",
    )
    add_caption(doc, "Таблица 1 — Стадии и сроки")
    add_table(
        doc,
        ["Стадия", "Работы", "Срок"],
        [
            ["Техническое задание", "Постановка задачи, разработка и утверждение ТЗ", "сентябрь 2026"],
            ["Рабочий проект", "Программа, документация, испытания по ПМИ", "октябрь–ноябрь 2026"],
            ["Внедрение", "Передача программы и документов преподавателю", "ноябрь 2026"],
        ],
        col_widths=[4.2, 8.6, 4.2],
        header_size=11,
        cell_size=11,
    )

    add_section_heading(doc, "8. Порядок контроля и приёмки")
    add_body(
        doc,
        "Виды испытаний: предварительные (исполнитель) и приёмо-сдаточные "
        "(в присутствии преподавателя). Контрольный пример: не менее "
        "12 товаров трёх категорий, 5 продаж (включая отказ без отметки "
        "18+ и продажу с доставкой в закрытой упаковке); полный цикл — "
        "приход, продажа, уменьшение остатка, приватный чек, отчёт "
        "«Остатки» и резервная копия. Приёмка — при соответствии "
        "программы требованиям настоящего ТЗ, наличии комплекта "
        "документации и успешном запуске на рабочем месте п. 4.4. "
        "Отказ в приёмке: невозможность запуска, продажа без отметки 18+, "
        "продажа сверх остатка, потеря данных при сохранении.",
    )

    add_section_heading(doc, "9. Приложение")
    add_center(doc, "(обязательное)", size=14, italic=True, space_after=4)
    add_center(doc, "Перечень терминов и нормативных документов", size=14, bold=True, space_after=8)
    add_body(
        doc,
        "ЕСПД — единая система программной документации; ТЗ — техническое "
        "задание; ПМИ — программа и методика испытаний; МТВ — настоящая "
        "программа; товар 18+ — позиция ассортимента, продаваемая только "
        "совершеннолетним; приватный чек — печатная форма без детализации "
        "наименований; остаток — количество, доступное к продаже.",
    )
    add_body(
        doc,
        "При разработке руководствуются: ГОСТ 19.101-77, ГОСТ 19.102-77, "
        "ГОСТ 19.104-78, ГОСТ 19.106-78, ГОСТ 19.201-78, ГОСТ 19.505-79, "
        "ГОСТ 19.301-79, ГОСТ 2.301-68. Продажа несовершеннолетним "
        "не допускается.",
    )

    def add_sign_line(text, space_before=0, keep_next=True):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.keep_with_next = keep_next
        run = p.add_run(text)
        set_run_font(run, size=14)

    add_sign_line(f"Исполнитель: {author['role']}", space_before=12, keep_next=True)
    add_sign_line(author["full"], keep_next=True)
    add_sign_line("«___» ______________ 2026 г.", keep_next=False)


def set_update_fields_on_open(doc):
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def build_document(author):
    global DOC_CODE
    DOC_CODE = author["code"]
    doc = Document()
    configure_normal_style(doc)
    setup_section(doc.sections[0], different_first=True)
    set_update_fields_on_open(doc)
    build_title_page(doc, author)
    doc.add_page_break()
    build_body(doc, author)
    doc.save(author["out"])
    print(author["out"])


def main():
    for author in STUDENTS:
        build_document(author)


if __name__ == "__main__":
    main()
